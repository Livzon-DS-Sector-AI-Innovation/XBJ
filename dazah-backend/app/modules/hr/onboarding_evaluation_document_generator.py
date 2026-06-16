"""员工上岗评估表 文档生成器."""

from datetime import date
from io import BytesIO
from pathlib import Path
from typing import Any

from openpyxl import Workbook, load_workbook
from openpyxl.styles import Alignment, Border, Side, Font
from docx import Document

from app.modules.hr.models import Employee
from app.modules.hr.schemas import OnboardingEvaluationInput


def _set_cell_format(cell, text: str) -> None:
    """设置单元格内容：宋体小四、居中."""
    from docx.oxml import parse_xml
    from docx.oxml.ns import nsdecls, qn

    tc = cell._tc
    for p in tc.findall(qn('w:p')):
        tc.remove(p)

    escaped = text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
    p_xml = (
        f'<w:p {nsdecls("w")}>'
        f'<w:pPr><w:jc w:val="center"/></w:pPr>'
        f'<w:r>'
        f'<w:rPr>'
        f'<w:rFonts w:ascii="宋体" w:hAnsi="宋体" w:eastAsia="宋体"/>'
        f'<w:sz w:val="24"/>'
        f'<w:szCs w:val="24"/>'
        f'</w:rPr>'
        f'<w:t>{escaped}</w:t>'
        f'</w:r>'
        f'</w:p>'
    )
    p = parse_xml(p_xml)
    tc.append(p)


def _fmt_date(value) -> str:
    """将日期格式化为 YYYY.MM.DD."""
    if not value:
        return ""
    if hasattr(value, "strftime"):
        return value.strftime("%Y.%m.%d")
    if isinstance(value, str):
        from datetime import datetime
        for fmt in ("%Y-%m-%d", "%Y/%m/%d", "%Y.%m.%d"):
            try:
                return datetime.strptime(value, fmt).strftime("%Y.%m.%d")
            except ValueError:
                continue
    return str(value)

NEW_TEMPLATE_NAME = "R-GN-2002 E 员工上岗评估表.docx"


def _find_new_template() -> Path:
    """Locate the new factory docx template."""
    candidates = [
        Path("新厂人员培训管理规程") / NEW_TEMPLATE_NAME,
        Path("../新厂人员培训管理规程") / NEW_TEMPLATE_NAME,
        Path(__file__).resolve().parent.parent.parent.parent
        / "新厂人员培训管理规程"
        / NEW_TEMPLATE_NAME,
    ]
    for p in candidates:
        if p.exists():
            return p
    raise FileNotFoundError(f"模板文件未找到: {NEW_TEMPLATE_NAME}")


def _cell_border():
    thin = Side(style="thin")
    return Border(left=thin, right=thin, top=thin, bottom=thin)


def _center_align():
    return Alignment(horizontal="center", vertical="center", wrap_text=True)


def _left_align():
    return Alignment(horizontal="left", vertical="center", wrap_text=True)


OLD_TEMPLATE_NAME = "7.12员工上岗评估表.xlsx"


def _find_old_template() -> Path:
    """Locate the old factory xlsx template."""
    candidates = [
        Path("员工培训教育管理规程") / OLD_TEMPLATE_NAME,
        Path("../员工培训教育管理规程") / OLD_TEMPLATE_NAME,
        Path(__file__).resolve().parent.parent.parent.parent
        / "员工培训教育管理规程"
        / OLD_TEMPLATE_NAME,
    ]
    for p in candidates:
        if p.exists():
            return p
    raise FileNotFoundError(f"模板文件未找到: {OLD_TEMPLATE_NAME}")


def _generate_old(data: OnboardingEvaluationInput) -> BytesIO:
    """Fill the old factory onboarding evaluation xlsx template."""
    try:
        template_path = _find_old_template()
        wb = load_workbook(str(template_path))
        ws = wb.active

        _data_font = Font(name="宋体", size=12)
        _data_align = Alignment(horizontal="center", vertical="center", wrap_text=True)

        # R4: 姓名 / 性别 / 所在部门/岗位
        ws["B4"] = data.employee_name or ""
        ws["B4"].font = _data_font
        ws["B4"].alignment = _data_align
        ws["D4"] = data.gender or ""
        ws["D4"].font = _data_font
        ws["D4"].alignment = _data_align
        ws["F4"] = data.department_position or ""
        ws["F4"].font = _data_font
        ws["F4"].alignment = _data_align

        # R5: 入厂时间 / 培训/考核期 / 转正时间
        ws["B5"] = _fmt_date(data.hire_date)
        ws["B5"].font = _data_font
        ws["B5"].alignment = _data_align
        ws["D5"] = data.training_period or ""
        ws["D5"].font = _data_font
        ws["D5"].alignment = _data_align
        ws["F5"] = _fmt_date(data.regularization_date)
        ws["F5"].font = _data_font
        ws["F5"].alignment = _data_align

        # R8-R13: 考核内容填写区域（6行）
        for i in range(6):
            row = 8 + i
            content_val = data.assessment_contents[i] if i < len(data.assessment_contents) else ""
            cell = ws[f"A{row}"]
            cell.value = content_val
            cell.font = _data_font
            cell.alignment = _data_align

        # R14: 评语
        ws["A14"] = data.comprehensive_comment or ""
        ws["A14"].font = _data_font
        ws["A14"].alignment = _data_align

        # R15: 同意上岗
        position = data.assigned_position or "____"
        agree_str = "☑" if data.is_qualified is True else "□"
        ws["A15"] = f" {agree_str}经考核该员工培训期表现优秀/确认，同意该员工正式上岗，担任{position}岗位。"
        ws["A15"].font = _data_font
        ws["A15"].alignment = _data_align

        # R16: 不同意上岗
        disagree_str = "☑" if data.is_qualified is False else "□"
        ws["A16"] = f" {disagree_str}经考核该员工培训期内表现不符合此岗位要求，不准上岗。"
        ws["A16"].font = _data_font
        ws["A16"].alignment = _data_align

        # R17: 考核方式
        method_map = {
            "理论": "☑理论 □实操 □现场",
            "实操": "□理论 ☑实操 □现场",
            "现场": "□理论 □实操 ☑现场",
        }
        method_str = method_map.get(data.assessment_method, "□理论 □实操 □现场")
        ws["A17"] = f" 考核方式：{method_str}"
        ws["A17"].font = _data_font
        ws["A17"].alignment = _data_align

        # R18: 部门负责人签名 / 日期
        sig_date = _fmt_date(data.signature_date)
        ws["A18"] = f" 部门负责人签名：{data.dept_manager_signature or ''}                   日期：{sig_date}"
        ws["A18"].font = _data_font
        ws["A18"].alignment = _data_align

        # R19: 备注
        ws["A19"] = f" 备注：{data.remarks or '培训期延长或转岗，由部门主管决定。'}"
        ws["A19"].font = _data_font
        ws["A19"].alignment = _data_align

        # R21-R23: 审批行
        approvals = [
            ("D21", "F21", data.dept_manager, data.dept_manager_agree),
            ("D22", "F22", data.hr_manager, data.hr_manager_agree),
            ("D23", "F23", data.qa_manager, data.qa_manager_agree),
        ]
        for name_cell, date_cell, name, agree in approvals:
            ws[name_cell] = name or ""
            ws[name_cell].font = _data_font
            ws[name_cell].alignment = _data_align
            ws[date_cell] = _fmt_date(data.approval_date)
            ws[date_cell].font = _data_font
            ws[date_cell].alignment = _data_align

        buffer = BytesIO()
        wb.save(buffer)
        buffer.seek(0)
        return buffer
    except Exception as e:
        import traceback
        with open('d:/LivzonAI/error.log', 'w', encoding='utf-8') as f:
            f.write(f"Error: {e}\n")
            f.write(traceback.format_exc())
        raise


def _generate_new(employee: Employee) -> BytesIO:
    """Fill the new factory onboarding evaluation docx template."""
    template_path = _find_new_template()
    doc = Document(str(template_path))

    if not doc.tables:
        raise ValueError("模板中未找到表格")

    table = doc.tables[0]

    # Row 0: [姓名] [] [部门] [] [工作卡号] []
    _set_cell_format(table.rows[0].cells[1], employee.name or "")
    _set_cell_format(table.rows[0].cells[3], employee.department or "")
    _set_cell_format(table.rows[0].cells[5], employee.employee_number or "")

    # Row 1: [学历] [] [专业] [] [毕业时间] []
    _set_cell_format(table.rows[1].cells[1], employee.education or "")
    _set_cell_format(table.rows[1].cells[3], employee.major or "")
    grad_date_str = _fmt_date(employee.graduation_date)
    _set_cell_format(table.rows[1].cells[5], grad_date_str)

    # Row 2: [报到日期] [] [考核人] [] [考核时间] []
    hire_date_str = _fmt_date(employee.hire_date)
    _set_cell_format(table.rows[2].cells[1], hire_date_str)
    _set_cell_format(table.rows[2].cells[3], "")
    _set_cell_format(table.rows[2].cells[5], "")

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def generate_onboarding_evaluation(data: Any, factory: str = "old") -> BytesIO:
    """根据员工数据或填写的评估信息生成上岗评估表文档.

    Returns a BytesIO buffer containing the generated document.
    """
    if factory == "new":
        if isinstance(data, Employee):
            return _generate_new(data)
        raise ValueError("新厂上岗评估表需要 Employee 对象")
    # 旧厂：如果传入 Employee，自动构建 OnboardingEvaluationInput
    if isinstance(data, Employee):
        input_data = OnboardingEvaluationInput(
            employee_name=data.name or "",
            gender=data.gender or "",
            department_position=f"{data.department or ''}/{data.position or ''}" if data.position else (data.department or ""),
            employee_number=data.employee_number or "",
            hire_date=data.hire_date,
        )
        return _generate_old(input_data)
    return _generate_old(data)
