"""Generate pre-job training plan documents from templates."""

from io import BytesIO
from pathlib import Path

import openpyxl
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from openpyxl.styles import Alignment, Font

from app.modules.hr.models import Employee


def _set_cell_format(cell, text: str) -> None:
    """设置单元格内容：宋体小四、居中。"""
    from docx.oxml import parse_xml
    from docx.oxml.ns import nsdecls, qn

    tc = cell._tc
    for p in tc.findall(qn('w:p')):
        tc.remove(p)

    escaped = text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
    p_xml = (
        f'<w:p {nsdecls("w")}>'
        f'<w:pPr><w:jc w:val="center"/></w:pPr>'
        f'<w:r>'
        f'<w:rPr>'
        f'<w:rFonts w:ascii="宋体" w:hAnsi="宋体" w:eastAsia="宋体"/>'
        f'<w:sz w:val="24"/>'
        f'<w:szCs w:val="24"/>'
        f'</w:rPr>'
        f'<w:t>{escaped}</w:t>'
        f'</w:r>'
        f'</w:p>'
    )
    p = parse_xml(p_xml)
    tc.append(p)


def _fmt_date(value) -> str:
    """将日期格式化为 YYYY.MM.DD."""
    if not value:
        return ""
    if hasattr(value, "strftime"):
        return value.strftime("%Y.%m.%d")
    if isinstance(value, str):
        from datetime import datetime
        for fmt in ("%Y-%m-%d", "%Y/%m/%d", "%Y.%m.%d"):
            try:
                return datetime.strptime(value, fmt).strftime("%Y.%m.%d")
            except ValueError:
                continue
    return str(value)


OLD_TEMPLATE_NAME = "7.4岗前培训计划.xlsx"
NEW_TEMPLATE_NAME = "R-GN-2002 C 岗前培训计划.docx"


def _find_old_template() -> Path:
    """Locate the old factory xlsx template."""
    candidates = [
        Path("员工培训教育管理规程") / OLD_TEMPLATE_NAME,
        Path("../员工培训教育管理规程") / OLD_TEMPLATE_NAME,
        Path(__file__).resolve().parent.parent.parent.parent
        / "员工培训教育管理规程"
        / OLD_TEMPLATE_NAME,
    ]
    for p in candidates:
        if p.exists():
            return p
    raise FileNotFoundError(f"模板文件未找到: {OLD_TEMPLATE_NAME}")


def _find_new_template() -> Path:
    """Locate the new factory docx template."""
    candidates = [
        Path("新厂人员培训管理规程") / NEW_TEMPLATE_NAME,
        Path("../新厂人员培训管理规程") / NEW_TEMPLATE_NAME,
        Path(__file__).resolve().parent.parent.parent.parent
        / "新厂人员培训管理规程"
        / NEW_TEMPLATE_NAME,
    ]
    for p in candidates:
        if p.exists():
            return p
    raise FileNotFoundError(f"模板文件未找到: {NEW_TEMPLATE_NAME}")


DEPT_CONTENT_MAP: dict[str, list[str]] = {
    "人事行政部": [
        "公司级公用文件(详见附件一)",
        "部门级公用文件(详见附件二)",
        "人事行政部人事行政专员岗位文件(详见附件三)",
        "人事行政专员岗位职责(QP.PM.053)",
        "生产安全知识",
        "岗前培训计划",
    ],
}


def _generate_old(employee: Employee) -> BytesIO:
    """Fill the old factory pre-job training plan xlsx template."""
    template_path = _find_old_template()
    wb = openpyxl.load_workbook(str(template_path))
    ws = wb.active

    _center = Alignment(horizontal="center", vertical="center")
    _font = Font(name="宋体", size=12)

    # Part 1: Employee overview
    ws["C5"] = employee.name or ""
    ws["C5"].font = _font
    ws["C5"].alignment = _center

    ws["I5"] = employee.department or ""
    ws["I5"].font = _font
    ws["I5"].alignment = _center

    ws["C6"] = employee.employee_number or ""
    ws["C6"].font = _font
    ws["C6"].alignment = _center

    hire_date_str = _fmt_date(employee.hire_date)
    ws["I6"] = hire_date_str
    ws["I6"].font = _font
    ws["I6"].alignment = _center

    ws["C7"] = employee.position or ""
    ws["C7"].font = _font
    ws["C7"].alignment = _center

    # Part 2: Training content (auto-fill by department)
    content_list = DEPT_CONTENT_MAP.get(employee.department or "", [])
    for i, content in enumerate(content_list):
        row = 11 + i
        if row <= 20:
            cell = ws[f"B{row}"]
            cell.value = content
            cell.font = _font
            cell.alignment = _center

    buffer = BytesIO()
    wb.save(buffer)
    buffer.seek(0)
    return buffer


def _generate_new(employee: Employee) -> BytesIO:
    """Fill the new factory pre-job training plan docx template."""
    template_path = _find_new_template()
    doc = Document(str(template_path))

    if not doc.tables:
        raise ValueError("模板中未找到表格")

    table = doc.tables[0]

    # Row 1: [姓名] [姓名] [] [] [部门] []
    # cells[0] and cells[1] are merged (label), cells[2] and cells[3] are merged (value), cells[4] is label, cells[5] is value
    _set_cell_format(table.rows[1].cells[2], employee.name or "")
    _set_cell_format(table.rows[1].cells[5], employee.department or "")

    # Row 2: [学历] [学历] [] [] [毕业院校] []
    _set_cell_format(table.rows[2].cells[2], employee.education or "")
    _set_cell_format(table.rows[2].cells[5], employee.school or "")

    # Row 3: [毕业时间] [毕业时间] [] [] [工作卡号] []
    grad_date_str = _fmt_date(employee.graduation_date)
    _set_cell_format(table.rows[3].cells[2], grad_date_str)
    _set_cell_format(table.rows[3].cells[5], employee.employee_number or "")

    # Row 4: [报到日期] [报到日期] [] [] [拟定岗位] []
    hire_date_str = _fmt_date(employee.hire_date)
    _set_cell_format(table.rows[4].cells[2], hire_date_str)
    _set_cell_format(table.rows[4].cells[5], employee.position or "")

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def generate_prejob_training_plan(employee: Employee, factory: str = "old") -> BytesIO:
    """Fill the pre-job training plan template with employee data.

    Returns a BytesIO buffer containing the generated document.
    """
    if factory == "new":
        return _generate_new(employee)
    return _generate_old(employee)
