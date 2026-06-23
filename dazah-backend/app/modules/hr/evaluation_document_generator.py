"""培训效果评估表 文档生成器."""

from datetime import date
from io import BytesIO
from pathlib import Path

from openpyxl import Workbook
from openpyxl.styles import Alignment, Border, Side, Font
from pydantic import BaseModel, Field
from docx import Document

def _fmt_date(value) -> str:
    """将日期格式化为 YYYY.MM.DD."""
    if not value:
        return ""
    if hasattr(value, "strftime"):
        return value.strftime("%Y.%m.%d")
    if isinstance(value, str):
        from datetime import datetime
        for fmt in ("%Y-%m-%d", "%Y/%m/%d", "%Y.%m.%d"):
            try:
                return datetime.strptime(value, fmt).strftime("%Y.%m.%d")
            except ValueError:
                continue
    return str(value)


def _set_cell_format(cell, text: str) -> None:
    """设置单元格内容：宋体小四、居中。"""
    from docx.oxml import parse_xml
    from docx.oxml.ns import nsdecls, qn

    tc = cell._tc
    for p in tc.findall(qn('w:p')):
        tc.remove(p)

    escaped = text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
    p_xml = (
        f'<w:p {nsdecls("w")}>'
        f'<w:pPr><w:jc w:val="center"/></w:pPr>'
        f'<w:r>'
        f'<w:rPr>'
        f'<w:rFonts w:ascii="宋体" w:hAnsi="宋体" w:eastAsia="宋体"/>'
        f'<w:sz w:val="24"/>'
        f'<w:szCs w:val="24"/>'
        f'</w:rPr>'
        f'<w:t>{escaped}</w:t>'
        f'</w:r>'
        f'</w:p>'
    )
    p = parse_xml(p_xml)
    tc.append(p)


NEW_TEMPLATE_NAME = "R-GN-2002 H 培训效果评估表.docx"
OLD_TEMPLATE_NAME = "7.11培训效果评估表.xlsx"


def _find_new_template() -> Path:
    """Locate the new factory docx template."""
    candidates = [
        Path("新厂人员培训管理规程") / NEW_TEMPLATE_NAME,
        Path("../新厂人员培训管理规程") / NEW_TEMPLATE_NAME,
        Path(__file__).resolve().parent.parent.parent.parent
        / "新厂人员培训管理规程"
        / NEW_TEMPLATE_NAME,
    ]
    for p in candidates:
        if p.exists():
            return p
    raise FileNotFoundError(f"模板文件未找到: {NEW_TEMPLATE_NAME}")


def _find_old_template() -> Path:
    """Locate the old factory xls template."""
    candidates = [
        Path("员工培训教育管理规程") / OLD_TEMPLATE_NAME,
        Path("../员工培训教育管理规程") / OLD_TEMPLATE_NAME,
        Path(__file__).resolve().parent.parent.parent.parent
        / "员工培训教育管理规程"
        / OLD_TEMPLATE_NAME,
    ]
    for p in candidates:
        if p.exists():
            return p
    raise FileNotFoundError(f"模板文件未找到: {OLD_TEMPLATE_NAME}")



from datetime import date
from io import BytesIO

from openpyxl import Workbook
from openpyxl.styles import Alignment, Border, Side, Font
from pydantic import BaseModel, Field


class TrainingEvaluationInput(BaseModel):
    subject: str = Field(..., max_length=128, description="培训主题")
    training_date: date | None = Field(None, description="培训日期")
    training_time_start: str | None = Field(None, max_length=32, description="培训开始时间")
    training_time_end: str | None = Field(None, max_length=32, description="培训结束时间")
    duration_hours: float | None = Field(None, description="学时")
    training_method: str | None = Field(None, max_length=32, description="培训方式")
    is_exam: bool = Field(False, description="是否考试")
    trainer_type: str | None = Field(None, max_length=64, description="培训人员类型")
    trainer: str | None = Field(None, max_length=64, description="授课人")
    department_personnel: str | None = Field(None, max_length=256, description="部门/班组/人员")
    expected_count: int | None = Field(None, description="应出席人数")
    actual_count: int | None = Field(None, description="实际出席人数")
    absent_count: int | None = Field(None, description="缺席人数")
    textbook: str | None = Field(None, max_length=256, description="培训教材")
    makeup_training: bool | None = Field(None, description="是否补课")
    assessment_method: str | None = Field(None, max_length=32, description="考核方式")
    pass_count: int | None = Field(None, description="合格人数")
    fail_count: int | None = Field(None, description="不合格人数")
    absent_exam_count: int | None = Field(None, description="缺考人数")
    absent_exam_handling: str | None = Field(None, max_length=512, description="缺考人员处理方式和原因")
    excellent_count: int | None = Field(None, description="优秀人数")
    qualified_count: int | None = Field(None, description="合格人数")
    unqualified_count: int | None = Field(None, description="不合格人数")
    evaluation_conclusion: str | None = Field(None, max_length=1024, description="培训效果评估及结论")
    organizer: str | None = Field(None, max_length=64, description="培训组织人")
    organizer_date: date | None = Field(None, description="组织日期")
    remarks: str | None = Field(None, max_length=512, description="备注")


def _cell_border():
    thin = Side(style="thin")
    return Border(left=thin, right=thin, top=thin, bottom=thin)


def _center_align():
    return Alignment(horizontal="center", vertical="center", wrap_text=True)


def _left_align():
    return Alignment(horizontal="left", vertical="center", wrap_text=True)


def _generate_old(data: TrainingEvaluationInput) -> BytesIO:
    """基于旧厂模板 7.11培训效果评估表.xlsx 填入数据，保留原始格式."""
    from openpyxl import load_workbook

    template_path = _find_old_template()
    wb = load_workbook(str(template_path))
    ws = wb.active

    # 培训主题：row 4, col 2 (B4)，合并区域 B4:E4
    ws.cell(row=4, column=2).value = data.subject or ""

    # 培训时间：row 5, col 2 (B5)，合并区域 B5:C5
    time_str = ""
    if data.training_date:
        time_str = data.training_date.strftime("%Y.%m.%d")
    if data.training_time_start and data.training_time_end:
        time_str += f" {data.training_time_start}~{data.training_time_end}"
    ws.cell(row=5, column=2).value = time_str

    # 学时：row 5, col 5 (E5)，D5 模板已有"学时"标签
    ws.cell(row=5, column=5).value = data.duration_hours

    # 培训方式：row 6, col 2 (B6)，合并区域 B6:C6，模板已有文字
    method_map = {
        "面授": "☑面授  □函授  □远程教育\n□自学  □其他",
        "函授": "□面授  ☑函授  □远程教育\n□自学  □其他",
        "远程教育": "□面授  □函授  ☑远程教育\n□自学  □其他",
        "自学": "□面授  □函授  □远程教育\n☑自学  □其他",
    }
    method_str = method_map.get(data.training_method, "□面授  □函授  □远程教育\n□自学  □其他")
    ws.cell(row=6, column=2).value = method_str

    # 授课人：row 6, col 5 (E6)，D6 模板已有"授课人"标签
    ws.cell(row=6, column=5).value = data.trainer or ""

    # 培训人员：C7 填入部门/班组/人员（B7 模板已有标签"部门/班组/人员"）
    ws.cell(row=7, column=3).value = data.textbook or ""

    # 应到/实到/缺席：row 8, col 2 (B8)，合并区域 B8:E8
    expected = data.expected_count if data.expected_count is not None else "___"
    actual = data.actual_count if data.actual_count is not None else "___"
    absent = data.absent_count if data.absent_count is not None else "___"
    ws.cell(row=8, column=2).value = f"应到 {expected} 人； 实到 {actual} 人； 缺席 {absent} 人。"

    # 培训教材：row 9, col 2 (B9)，合并区域 B9:E9，填培训内容
    ws.cell(row=9, column=2).value = data.subject or ""

    # 缺席人员处理方式：row 10, col 1 (A10)，合并区域 A10:E10
    makeup_str = ""
    if data.makeup_training is True:
        makeup_str = "是否进行补课培训，☑是  □否，未参加培训人员必须补上培训内容，（包括培训时间、地点、方式等）。"
    elif data.makeup_training is False:
        makeup_str = "是否进行补课培训，□是  ☑否，未参加培训人员必须补上培训内容，（包括培训时间、地点、方式等）。"
    else:
        makeup_str = "是否进行补课培训，□是  □否，未参加培训人员必须补上培训内容，（包括培训时间、地点、方式等）。"
    ws.cell(row=10, column=1).value = f"缺席人员处理方式：\n\n{makeup_str}\n"

    # 考核方式：row 11, col 2 (B11)，合并区域 B11:E11
    am = data.assessment_method or ""
    am_map = {
        "笔试": "☑ 笔试    □ 口试   □ 实操   □ 写总结  ",
        "口试": "□ 笔试    ☑ 口试   □ 实操   □ 写总结  ",
        "实操": "□ 笔试    □ 口试   ☑ 实操   □ 写总结  ",
        "写总结": "□ 笔试    □ 口试   □ 实操   ☑ 写总结  ",
    }
    ws.cell(row=11, column=2).value = am_map.get(am, "□ 笔试    □ 口试   □ 实操   □ 写总结  ")

    # 考核结果：row 12, col 2 (B12)，合并区域 B12:E12
    p_cnt = data.pass_count if data.pass_count is not None else "___"
    f_cnt = data.fail_count if data.fail_count is not None else "___"
    ae_cnt = data.absent_exam_count if data.absent_exam_count is not None else "___"
    ws.cell(row=12, column=2).value = f"合格 {p_cnt} 人；不合格 {f_cnt} 人；缺考 {ae_cnt} 人。"

    # 缺考人员处理方式：row 13, col 1 (A13)，合并区域 A13:E13
    ws.cell(row=13, column=1).value = f"缺考人员处理方式和原因：{data.absent_exam_handling or ''}"

    # 综合评分：row 14, col 2 (B14)，合并区域 B14:E14
    ex_cnt = data.excellent_count if data.excellent_count is not None else "___"
    q_cnt = data.qualified_count if data.qualified_count is not None else "___"
    uq_cnt = data.unqualified_count if data.unqualified_count is not None else "___"
    ws.cell(row=14, column=2).value = f"优秀 {ex_cnt} 人；合格 {q_cnt} 人；不合格 {uq_cnt} 人。"

    # 缺考/不合格人员处理方式：row 15, col 1 (A15)，合并区域 A15:E15
    ws.cell(row=15, column=1).value = "缺考/不合格人员处理方式："

    # 培训效果评估及结论：row 16, col 1 (A16)，合并区域 A16:E19
    ws.cell(row=16, column=1).value = f"培训效果评估及结论：\n{data.evaluation_conclusion or ''}"

    # 培训组织人/日期：row 20, col 1 (A20)，合并区域 A20:E20
    org_str = ""
    if data.organizer:
        org_str = f"                                          {data.organizer}"
    if data.organizer_date:
        org_str += f" / {data.organizer_date.strftime('%Y.%m.%d')}"
    ws.cell(row=20, column=1).value = f"培训组织人/日期：{org_str}"

    # 备注：row 21, col 1 (A21)，合并区域 A21:E21
    ws.cell(row=21, column=1).value = f"备注：\n{data.remarks or ''}"

    buffer = BytesIO()
    wb.save(buffer)
    buffer.seek(0)
    return buffer


def _generate_new(data: TrainingEvaluationInput) -> BytesIO:
    """Fill the new factory training evaluation docx template."""
    template_path = _find_new_template()
    doc = Document(str(template_path))

    if not doc.tables:
        raise ValueError("模板中未找到表格")

    table = doc.tables[0]

    # Row 0: 培训内容（全行合并）→ 填入主题
    topic = data.subject or ""
    _set_cell_format(table.rows[0].cells[0], f"培训内容：{topic}")

    # Row 1: 培训日期(col1) | 课时(col3)
    _set_cell_format(table.rows[1].cells[1], _fmt_date(data.training_date))
    _set_cell_format(table.rows[1].cells[3], str(data.duration_hours) if data.duration_hours is not None else "")

    # Row 2: 培训方式(col1 checkbox) | 授课人(col3)
    method_str = data.training_method or ""
    _set_cell_format(table.rows[2].cells[1], method_str)
    _set_cell_format(table.rows[2].cells[3], data.trainer or "")

    # Row 3: 培训教材 → 填入右边的格子里
    _set_cell_format(table.rows[3].cells[1], data.subject or "")

    # Row 4: 培训对象 → 部门/班组/人员
    _set_cell_format(table.rows[4].cells[1], f"部门/班组/人员：{data.textbook or ''}")

    # Row 5: 培训对象 → 应到/实到/缺席
    expected = data.expected_count if data.expected_count is not None else ""
    actual = data.actual_count if data.actual_count is not None else ""
    absent = data.absent_count if data.absent_count is not None else ""
    _set_cell_format(table.rows[5].cells[1], f"应到：{expected} 人， 实到：{actual} 人， 缺席：{absent} 人。")

    # Row 6: 缺席人员处理方式（全行合并）
    makeup_str = ""
    if data.makeup_training is True:
        makeup_str = "是否进行再培训 ☑否  □是 （若再培训，请填写以下资料）\n再培训（时间、地点、方式等）："
    elif data.makeup_training is False:
        makeup_str = "是否进行再培训 □否  ☑是 （若再培训，请填写以下资料）\n再培训（时间、地点、方式等）："
    else:
        makeup_str = "是否进行再培训 □否  □是 （若再培训，请填写以下资料）\n再培训（时间、地点、方式等）："
    _set_cell_format(table.rows[6].cells[0], f"缺席人员处理方式：\n\n{makeup_str}")

    # Row 7: 考核方式（cols1-3合并）
    am = data.assessment_method or ""
    _set_cell_format(table.rows[7].cells[1], am)

    # Row 8: 考核结果（cols1-3合并）
    p_cnt = data.pass_count if data.pass_count is not None else ""
    f_cnt = data.fail_count if data.fail_count is not None else ""
    ae_cnt = data.absent_exam_count if data.absent_exam_count is not None else ""
    _set_cell_format(table.rows[8].cells[1], f"合格：{p_cnt} 人；不合格：{f_cnt} 人；缺考：{ae_cnt} 人。")

    # Row 9: 缺考及不合格人员处理方式（全行合并）
    _set_cell_format(table.rows[9].cells[0], f"缺考及不合格人员处理方式：{data.absent_exam_handling or ''}")

    # Row 10: 补考结果（cols1-3合并）
    _set_cell_format(table.rows[10].cells[1], f"补考：    人；合格：    人；不合格：    人。")

    # Row 11: 缺考及补考不合格人员处理方式（全行合并）
    _set_cell_format(table.rows[11].cells[0], "缺考及补考不合格人员处理方式：")

    # Row 12: 培训效果评估及其他（全行合并，多行）
    conclusion = data.evaluation_conclusion or ""
    organizer = data.organizer or ""
    org_date = _fmt_date(data.organizer_date)
    _set_cell_format(table.rows[12].cells[0], f"培训效果评估及其他：\n{conclusion}\n\n\n\n\n\n\n\n\n\n培训考核人/日期：{organizer} / {org_date}")

    # Row 13: 其他（全行合并）
    _set_cell_format(table.rows[13].cells[0], f"其他：{data.remarks or ''}")

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def generate_training_evaluation(data: TrainingEvaluationInput, factory: str = "old") -> BytesIO:
    """根据填写的培训信息生成培训效果评估表文档."""
    if factory == "new":
        return _generate_new(data)
    return _generate_old(data)
