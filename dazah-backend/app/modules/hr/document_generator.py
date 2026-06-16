"""Generate onboarding training record documents from templates."""

from io import BytesIO
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt

from app.modules.hr.models import Employee


def _set_cell_format(cell, text: str) -> None:
    """设置单元格内容：宋体小四、居中."""
    from docx.oxml import parse_xml
    from docx.oxml.ns import nsdecls, qn

    tc = cell._tc
    for p in tc.findall(qn('w:p')):
        tc.remove(p)

    escaped = text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
    p_xml = (
        f'<w:p {nsdecls("w")}>'
        f'<w:pPr><w:jc w:val="center"/></w:pPr>'
        f'<w:r>'
        f'<w:rPr>'
        f'<w:rFonts w:ascii="宋体" w:hAnsi="宋体" w:eastAsia="宋体"/>'
        f'<w:sz w:val="24"/>'
        f'<w:szCs w:val="24"/>'
        f'</w:rPr>'
        f'<w:t>{escaped}</w:t>'
        f'</w:r>'
        f'</w:p>'
    )
    p = parse_xml(p_xml)
    tc.append(p)


def _fmt_date(value) -> str:
    """将日期格式化为 YYYY.MM.DD."""
    if not value:
        return ""
    if hasattr(value, "strftime"):
        return value.strftime("%Y.%m.%d")
    if isinstance(value, str):
        from datetime import datetime
        for fmt in ("%Y-%m-%d", "%Y/%m/%d", "%Y.%m.%d"):
            try:
                return datetime.strptime(value, fmt).strftime("%Y.%m.%d")
            except ValueError:
                continue
    return str(value)


OLD_TEMPLATE_NAME = "7.3新员工入职培训记录.docx"
NEW_TEMPLATE_NAME = "R-GN-2002 B 新员工入职培训记录.docx"


def _find_template(factory: str = "old") -> Path:
    """Locate the docx template, trying several path candidates."""
    if factory == "new":
        candidates = [
            Path("新厂人员培训管理规程") / NEW_TEMPLATE_NAME,
            Path("../新厂人员培训管理规程") / NEW_TEMPLATE_NAME,
            Path(__file__).resolve().parent.parent.parent.parent
            / "新厂人员培训管理规程"
            / NEW_TEMPLATE_NAME,
        ]
    else:
        candidates = [
            Path("员工培训教育管理规程") / OLD_TEMPLATE_NAME,
            Path("../员工培训教育管理规程") / OLD_TEMPLATE_NAME,
            Path(__file__).resolve().parent.parent.parent.parent
            / "员工培训教育管理规程"
            / OLD_TEMPLATE_NAME,
        ]
    for p in candidates:
        if p.exists():
            return p
    raise FileNotFoundError(f"模板文件未找到: {NEW_TEMPLATE_NAME if factory == 'new' else OLD_TEMPLATE_NAME}")


def generate_onboarding_training_record(employee: Employee, factory: str = "old") -> BytesIO:
    """Fill the onboarding training record template with employee data.

    Returns a BytesIO buffer containing the generated docx.
    """
    template_path = _find_template(factory)
    doc = Document(str(template_path))

    if not doc.tables:
        raise ValueError("模板中未找到表格")

    table = doc.tables[0]

    if factory == "new":
        # New factory template: 24 rows x 6 cols
        # Row 1: [姓名] [] [性别] [] [工作卡号] []
        _set_cell_format(table.rows[1].cells[1], employee.name or "")
        _set_cell_format(table.rows[1].cells[3], employee.gender or "")
        _set_cell_format(table.rows[1].cells[5], employee.employee_number or "")

        # Row 2: [部门] [] [] [拟定岗位] [] []
        _set_cell_format(table.rows[2].cells[1], employee.department or "")
        _set_cell_format(table.rows[2].cells[4], employee.position or "")

        # Row 3: [报到日期] [] [] [拟定转正日期] [] []
        hire_date_str = _fmt_date(employee.hire_date)
        _set_cell_format(table.rows[3].cells[1], hire_date_str)
        _set_cell_format(table.rows[3].cells[4], "")
    else:
        # Old factory template: 21 rows x 12 cols
        # Row 1: 姓名 | 01 | 性别 | 02... | 工作卡号 | 03...
        _set_cell_format(table.rows[1].cells[1], employee.name or "")
        for idx in (3, 4, 5):
            _set_cell_format(table.rows[1].cells[idx], employee.gender or "")
        for idx in (9, 10, 11):
            _set_cell_format(table.rows[1].cells[idx], employee.employee_number or "")

        # Row 2: 部门 | 04... | 拟定岗位 | 05...
        for idx in (1, 2):
            _set_cell_format(table.rows[2].cells[idx], employee.department or "")
        for idx in (6, 7, 8, 9, 10, 11):
            _set_cell_format(table.rows[2].cells[idx], employee.position or "")

        # Row 3: 报到日期 | 06... | 转正日期 | 07...
        hire_date_str = _fmt_date(employee.hire_date)
        for idx in (1, 2):
            _set_cell_format(table.rows[3].cells[idx], hire_date_str)
        # 转正日期数据库中无对应字段，留空
        for idx in (6, 7, 8, 9, 10, 11):
            _set_cell_format(table.rows[3].cells[idx], "")

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
